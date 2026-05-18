from __future__ import annotations

from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase
from unittest.mock import patch
import wave

from docx import Document
from PIL import Image
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.providers import TranscriptResult, TranscriptSegment
import app.core.session_assets as session_asset_utils
import app.services.asset_ingestion_service as asset_ingestion_service
import app.services.asset_service as asset_service
from app.db.base import Base
from app.models import SessionRecord


class AssetIngestionServiceTests(TestCase):
    def setUp(self) -> None:
        super().setUp()
        self.tempdir = TemporaryDirectory()
        self.engine = create_engine("sqlite:///:memory:", future=True)
        Base.metadata.create_all(bind=self.engine)
        self.SessionLocal = sessionmaker(bind=self.engine, autoflush=False, autocommit=False, future=True)
        self.session_id = "session-1"

        with self._create_session() as db:
            db.add(
                SessionRecord(
                    id=self.session_id,
                    title="Test Session",
                    workspace_mode="bound",
                    canonical_workspace_path="/tmp/workspace",
                    workspace_fingerprint="workspace-fp",
                    workspace_label="workspace",
                    status="idle",
                )
            )
            db.commit()

    def tearDown(self) -> None:
        self.tempdir.cleanup()
        super().tearDown()

    @contextmanager
    def _create_session(self):
        db = self.SessionLocal()
        try:
            yield db
        finally:
            db.close()

    def test_validate_upload_rejects_unsupported_type(self) -> None:
        with self.assertRaises(asset_ingestion_service.AssetUploadError):
            asset_ingestion_service.validate_upload("notes.zip", "application/zip", 128)

    def test_validate_upload_accepts_audio_and_video_types(self) -> None:
        self.assertEqual(asset_ingestion_service.validate_upload("voice.wav", "audio/wav", 1024), "audio")
        self.assertEqual(asset_ingestion_service.validate_upload("clip.mp4", "video/mp4", 1024), "video")

    def test_image_ingestion_creates_preview_and_marks_ready(self) -> None:
        image = Image.new("RGB", (640, 480), color=(120, 80, 10))
        image_bytes = BytesIO()
        image.save(image_bytes, format="PNG")

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "screen.png",
                "image/png",
                image_bytes.getvalue(),
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)

        self.assertEqual(ingested.status, "ready")
        self.assertTrue(ingested.preview_path)
        self.assertTrue(Path(ingested.preview_path).exists())

    def test_docx_ingestion_creates_chunks(self) -> None:
        document = Document()
        document.add_heading("Status", level=1)
        document.add_paragraph("The runtime is now asset aware.")
        document.add_paragraph("The composer will eventually support drag and drop.")
        docx_path = Path(self.tempdir.name) / "sample.docx"
        document.save(docx_path)

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "sample.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                docx_path.read_bytes(),
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)
            chunks = asset_service.list_asset_chunks(asset.id)

        self.assertEqual(ingested.status, "ready")
        self.assertGreaterEqual(len(chunks), 1)
        self.assertTrue(any("asset aware" in chunk.content for chunk in chunks))

    def test_stage_uploaded_asset_stream_writes_file_incrementally(self) -> None:
        class FakeUpload:
            def __init__(self, filename: str, content_type: str, chunks: list[bytes]) -> None:
                self.filename = filename
                self.content_type = content_type
                self._chunks = list(chunks)

            async def read(self, size: int = -1) -> bytes:
                return self._chunks.pop(0) if self._chunks else b""

        upload = FakeUpload("streamed.txt.pdf", "application/pdf", [b"%PDF-1.4\n", b"body", b"tail"])

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ):
            asset = self._run_async(
                asset_ingestion_service.stage_uploaded_asset_stream(self.session_id, upload, chunk_size=4)
            )

        self.assertEqual(asset.status, "uploaded")
        self.assertEqual(asset.size_bytes, len(b"%PDF-1.4\nbodytail"))
        self.assertTrue(Path(asset.storage_path).exists())

    def test_audio_ingestion_stores_basic_metadata(self) -> None:
        wav_path = Path(self.tempdir.name) / "sample.wav"
        with wave.open(wav_path.as_posix(), "wb") as handle:
            handle.setnchannels(1)
            handle.setsampwidth(2)
            handle.setframerate(8000)
            handle.writeframes(b"\x00\x00" * 800)

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ), patch.object(
            asset_ingestion_service.settings,
            "jarvis_speech_recognition_provider",
            "",
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "sample.wav",
                "audio/wav",
                wav_path.read_bytes(),
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)

        self.assertEqual(ingested.kind, "audio")
        self.assertEqual(ingested.status, "ready")
        self.assertEqual(ingested.metadata_json.get("container"), "wav")
        self.assertEqual(ingested.metadata_json.get("sample_rate"), 8000)
        self.assertEqual(ingested.metadata_json.get("channels"), 1)
        self.assertEqual(ingested.metadata_json.get("transcript_status"), "not_configured")

    def test_video_ingestion_stores_basic_metadata(self) -> None:
        video_bytes = b"\x00\x00\x00\x18ftypmp42\x00\x00\x00\x00mp42isom"

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ), patch.object(
            asset_ingestion_service.settings,
            "jarvis_speech_recognition_provider",
            "",
        ), patch(
            "app.services.asset_ingestion_service._extract_video_keyframes",
            return_value=("not_configured", []),
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "sample.mp4",
                "video/mp4",
                video_bytes,
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)

        self.assertEqual(ingested.kind, "video")
        self.assertEqual(ingested.status, "ready")
        self.assertEqual(ingested.metadata_json.get("container"), "mp4")
        self.assertEqual(ingested.metadata_json.get("transcript_status"), "not_configured")
        self.assertEqual(ingested.metadata_json.get("keyframe_status"), "not_configured")

    def test_video_ingestion_writes_transcript_chunks_when_asr_is_available(self) -> None:
        video_path = Path(self.tempdir.name) / "sample.mp4"
        video_path.write_bytes(b"\x00\x00\x00\x18ftypmp42\x00\x00\x00\x00mp42isom")
        extracted_audio_path = Path(self.tempdir.name) / "extracted-audio.wav"
        extracted_audio_path.write_bytes(b"RIFFxxxxWAVE")
        transcript = TranscriptResult(
            text="describe the plan",
            segments=[
                TranscriptSegment(text="describe the plan", start_ms=120, end_ms=1480),
            ],
            provider_name="fake-asr",
            metadata={"duration_ms": 1480},
        )

        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ), patch.object(
            asset_ingestion_service.settings,
            "jarvis_speech_recognition_provider",
            "volcengine",
        ), patch(
            "app.services.asset_ingestion_service._extract_video_audio_track",
            return_value=extracted_audio_path,
        ), patch(
            "app.services.asset_ingestion_service._extract_video_keyframes",
            return_value=("ready", [{"frame_index": 0, "timestamp_ms": 800, "path": str(Path(self.tempdir.name) / 'keyframe-000.png')}]),
        ), patch(
            "app.services.asset_ingestion_service.speech_recognition_service.transcribe",
            return_value=transcript,
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "sample.mp4",
                "video/mp4",
                video_path.read_bytes(),
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)
            chunks = asset_service.list_asset_chunks(asset.id)

        self.assertEqual(ingested.metadata_json.get("transcript_status"), "ready")
        self.assertEqual(ingested.metadata_json.get("derived_audio_path"), extracted_audio_path.as_posix())
        self.assertEqual(ingested.metadata_json.get("keyframe_status"), "ready")
        self.assertEqual(ingested.metadata_json.get("keyframe_count"), 1)
        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0].content, "describe the plan")
        self.assertEqual(chunks[0].start_ms, 120)
        self.assertEqual(chunks[0].end_ms, 1480)
        self.assertEqual(chunks[1].frame_index, 0)

    def test_audio_ingestion_writes_transcript_chunks_when_asr_is_available(self) -> None:
        wav_path = Path(self.tempdir.name) / "transcribe.wav"
        with wave.open(wav_path.as_posix(), "wb") as handle:
            handle.setnchannels(1)
            handle.setsampwidth(2)
            handle.setframerate(8000)
            handle.writeframes(b"\x00\x00" * 800)

        transcript = TranscriptResult(
            text="hello world",
            segments=[
                TranscriptSegment(text="hello", start_ms=0, end_ms=500),
                TranscriptSegment(text="world", start_ms=600, end_ms=1200),
            ],
            provider_name="fake-asr",
            metadata={"duration_ms": 1200},
        )
        with patch.object(asset_service, "create_session", self._create_session), patch.object(
            session_asset_utils.settings,
            "data_dir",
            Path(self.tempdir.name),
        ), patch.object(
            asset_ingestion_service.settings,
            "jarvis_speech_recognition_provider",
            "volcengine",
        ), patch(
            "app.services.asset_ingestion_service.speech_recognition_service.transcribe",
            return_value=transcript,
        ):
            asset = asset_ingestion_service.stage_uploaded_asset(
                self.session_id,
                "transcribe.wav",
                "audio/wav",
                wav_path.read_bytes(),
            )
            ingested = asset_ingestion_service.ingest_asset(asset.id)
            chunks = asset_service.list_asset_chunks(asset.id)

        self.assertEqual(ingested.metadata_json.get("transcript_status"), "ready")
        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0].content, "hello")
        self.assertEqual(chunks[0].start_ms, 0)
        self.assertEqual(chunks[1].end_ms, 1200)

    def _run_async(self, awaitable):
        import asyncio

        return asyncio.run(awaitable)
